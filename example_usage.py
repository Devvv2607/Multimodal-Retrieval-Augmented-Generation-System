"""
Example usage script for the multimodal RAG system.
"""

import os
import sys
import tempfile

# Add the current directory to Python path
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

def create_sample_data():
    """Create sample data files for demonstration."""
    # Create sample directory
    sample_dir = "./sample_data"
    os.makedirs(sample_dir, exist_ok=True)
    
    # Create sample text document
    text_content = """
    The field of artificial intelligence has seen tremendous growth in recent years. 
    Machine learning algorithms, particularly deep learning neural networks, have 
    achieved remarkable success in various domains including computer vision, 
    natural language processing, and speech recognition.
    
    One of the most significant breakthroughs has been the development of 
    transformer architectures, which have revolutionized natural language 
    understanding tasks. Models like BERT, GPT, and more recently Phi-3 have 
    demonstrated unprecedented capabilities in generating human-like text and 
    understanding context.
    
    In the realm of computer vision, convolutional neural networks (CNNs) have 
    become the standard approach for image classification, object detection, 
    and segmentation tasks. The integration of vision and language models has 
    led to multimodal systems that can understand both textual and visual 
    information simultaneously.
    """
    
    with open(os.path.join(sample_dir, "ai_overview.txt"), "w") as f:
        f.write(text_content)
    
    # Create sample DOCX document
    try:
        from docx import Document
        doc = Document()
        doc.add_heading('Machine Learning Fundamentals', 0)
        doc.add_paragraph(
            'Machine learning is a subset of artificial intelligence that focuses '
            'on developing algorithms that can learn from and make predictions or '
            'decisions based on data.'
        )
        doc.add_paragraph(
            'The three main types of machine learning are:\n'
            '1. Supervised Learning - Learning with labeled data\n'
            '2. Unsupervised Learning - Finding patterns in unlabeled data\n'
            '3. Reinforcement Learning - Learning through interaction with an environment'
        )
        doc.save(os.path.join(sample_dir, "ml_fundamentals.docx"))
    except ImportError:
        # If python-docx is not available, create a text file instead
        docx_content = """
        Machine Learning Fundamentals
        
        Machine learning is a subset of artificial intelligence that focuses 
        on developing algorithms that can learn from and make predictions or 
        decisions based on data.
        
        The three main types of machine learning are:
        1. Supervised Learning - Learning with labeled data
        2. Unsupervised Learning - Finding patterns in unlabeled data
        3. Reinforcement Learning - Learning through interaction with an environment
        """
        with open(os.path.join(sample_dir, "ml_fundamentals.txt"), "w") as f:
            f.write(docx_content)
    
    print(f"Created sample data in {sample_dir}")
    return sample_dir

def demonstrate_ingestion():
    """Demonstrate the ingestion process."""
    print("=== Demonstrating Ingestion Process ===")
    
    # Create sample data
    sample_dir = create_sample_data()
    
    # Import ingestor
    from ingestion.ingestor import Ingestor
    
    # Initialize ingestor
    ingestor = Ingestor()
    
    # Ingest sample data
    print(f"Ingesting data from {sample_dir}...")
    count = ingestor.ingest_directory(sample_dir)
    print(f"Successfully ingested {count} files")
    
    return True

def demonstrate_query():
    """Demonstrate the query process."""
    print("\n=== Demonstrating Query Process ===")
    
    # Import required components
    from retrieval.retriever import Retriever
    from generation.generator import Generator
    from indexing.vector_store import VectorStore
    from embedding.text_embedder import TextEmbedder
    from embedding.image_embedder import ImageEmbedder
    from utils.config import config
    
    # Initialize components
    vector_store = VectorStore(
        dimension=config.get('models.text_embedding.dim', 384),
        index_path=config.get('vector_db.index_path'),
        metadata_path=config.get('vector_db.metadata_path')
    )
    
    text_embedder = TextEmbedder(config.get('models.text_embedding.name'))
    image_embedder = ImageEmbedder(config.get('models.image_embedding.name'))
    retriever = Retriever(vector_store, text_embedder, image_embedder)
    generator = Generator(
        model_name=config.get('models.llm.name'),
        max_tokens=config.get('models.llm.max_tokens', 2048)
    )
    
    # Example queries
    queries = [
        "What are the main types of machine learning?",
        "How has transformer architecture impacted AI?",
        "What role do convolutional neural networks play in computer vision?"
    ]
    
    for query in queries:
        print(f"\nQuestion: {query}")
        
        # Retrieve context
        context = retriever.retrieve_text(query, k=3)
        
        if context:
            # Generate answer
            answer = generator.generate_answer(query, context)
            print(f"Answer: {answer}")
            
            # Show sources
            print("Sources:")
            for i, item in enumerate(context, 1):
                source = item.get('source', 'Unknown')
                print(f"  [{i}] {source}")
        else:
            print("No relevant context found.")
    
    return True

def main():
    """Main demonstration function."""
    print("Multimodal RAG System - Example Usage")
    print("=" * 40)
    
    try:
        # Initialize system
        from init import initialize_system
        initialize_system()
        
        # Demonstrate ingestion
        if demonstrate_ingestion():
            # Demonstrate query
            demonstrate_query()
        
        print("\n=== Demonstration Complete ===")
        print("You can now use the system with:")
        print("  python main.py ingest --input_dir ./sample_data")
        print("  python main.py query --question \"Your question here\"")
        print("  python main.py --gui (for graphical interface)")
        
    except Exception as e:
        print(f"Error during demonstration: {str(e)}")
        return False
    
    return True

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)